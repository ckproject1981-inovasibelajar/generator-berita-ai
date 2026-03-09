import streamlit as st
import google.generativeai as genai
from docx import Document
import io

# Konfigurasi Halaman
st.set_page_config(page_title="i-Humas DLI", layout="centered")

# --- CSS AGRESIF UNTUK CLEAN LOOK & FOOTER KUSTOM ---
hide_streamlit_style = """
    <style>
    #MainMenu {visibility: hidden !important;}
    footer {visibility: hidden !important;}
    .stDeployButton {display:none !important;}
    header {visibility: hidden !important;}
    button[title="View fullscreen"] {visibility: hidden !important;}
    .stAppDeployButton {display:none !important;}
    
    .custom-footer {
        position: fixed;
        left: 0;
        bottom: 0;
        width: 100%;
        text-align: center;
        padding: 10px;
        background-color: #f8f9fa;
        font-size: 12px;
        color: #555;
        border-top: 1px solid #ddd;
        z-index: 999;
    }
    </style>
"""
st.markdown(hide_streamlit_style, unsafe_allow_html=True)

# --- LOGO & JUDUL ---
# Menampilkan logo dengan lebar 200px dan rata tengah
col_logo1, col_logo2, col_logo3 = st.columns([1, 2, 1])
with col_logo2:
    st.image("https://i.ibb.co.com/23N3kpBY/Logo-DLI.png", width=200)

st.markdown("""
    <h2 style="font-size: 20px; text-align: center; margin-top: -10px;">📰 i-Humas PUI-PT DLI UM Ver 1.0</h2>
""", unsafe_allow_html=True)

# Logika API Key
if "GOOGLE_API_KEY" in st.secrets:
    api_key = st.secrets["GOOGLE_API_KEY"]
else:
    api_key = st.sidebar.text_input("Masukkan API Key (Lokal)", type="password")

if api_key:
    genai.configure(api_key=api_key)

# Form Input
with st.form("news_form"):
    judul_awal = st.text_input("Judul Draft (Opsional)")
    tempat = st.text_input("Tempat Kegiatan")
    waktu = st.text_input("Waktu Kegiatan")
    narasumber = st.text_input("Narasumber")
    gaya = st.selectbox("Gaya Bahasa", ["Formal", "Santai", "Investigatif", "Sensasional"])
    rincian = st.text_area("Rincian Kegiatan (Data Utama)", height=150, help="Masukkan fakta-fakta kegiatan di sini.")
    kata_kunci = st.text_input("Kata Kunci (pisahkan dengan koma)")
    submit = st.form_submit_button("Generate Berita Profesional")

# Fungsi untuk membuat file Word
def create_docx(text):
    doc = Document()
    doc.add_heading('Hasil Berita AI', 0)
    doc.add_paragraph(text)
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

if submit:
    if not api_key:
        st.error("API Key belum diisi!")
    else:
        try:
            model = genai.GenerativeModel('gemini-1.5-flash')
            
            # --- SARAN JUDUL ALTERNATIF ---
            saran_prompt = f"Berikan 3 saran judul berita yang menarik, profesional, dan SEO-friendly berdasarkan detail kegiatan ini: {rincian}. Berikan hanya judulnya saja."
            saran_response = model.generate_content(saran_prompt)
            st.info(f"💡 **Saran Judul Alternatif:**\n{saran_response.text}")

            # --- GENERATE BERITA 5W+1H ---
            prompt = f"""
            Anda adalah jurnalis senior i-Humas PUI-PT DLI UM. Tulis berita profesional berdasarkan:
            - Judul Draft: {judul_awal}
            - Lokasi: {tempat}
            - Waktu: {waktu}
            - Narasumber: {narasumber}
            - Rincian: {rincian}
            - Gaya: {gaya}
            
            INSTRUKSI: Terapkan 5W+1H, struktur berita piramida terbalik, bahasa objektif, dan informatif.
            """
            
            with st.spinner("Jurnalis AI sedang menulis berita..."):
                response = model.generate_content(prompt)
                berita_teks = response.text
                st.markdown("### Hasil Berita Final:")
                st.write(berita_teks)
                
                doc_buffer = create_docx(berita_teks)
                st.download_button("📥 Download Berita (.docx)", doc_buffer, "berita_dli_um.docx")
        except Exception as e:
            st.error(f"Terjadi kesalahan: {e}")

# --- FOOTER KUSTOM ---
st.markdown('<div class="custom-footer">Dikembangkan oleh Citra Kurniawan & PUI DLI - 2026</div>', unsafe_allow_html=True)